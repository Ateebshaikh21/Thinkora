import io
import logging
from typing import Optional, Tuple
from pathlib import Path

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class FileProcessor:
    """
    Utility class to extract text content from various file formats
    """
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv',
        '.md': 'text/markdown',
        '.rtf': 'application/rtf'
    }
    
    @classmethod
    def get_supported_extensions(cls) -> list:
        """Get list of supported file extensions"""
        return list(cls.SUPPORTED_EXTENSIONS.keys())
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file format is supported"""
        extension = Path(filename).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def get_file_type(cls, filename: str) -> str:
        """Get file type based on extension"""
        extension = Path(filename).suffix.lower()
        
        if extension in ['.pdf']:
            return 'pdf'
        elif extension in ['.docx', '.doc']:
            return 'word'
        elif extension in ['.xlsx', '.xls']:
            return 'excel'
        elif extension in ['.csv']:
            return 'csv'
        elif extension in ['.txt', '.md', '.rtf']:
            return 'text'
        else:
            return 'unknown'
    
    @classmethod
    async def extract_text(cls, file_content: bytes, filename: str) -> Tuple[str, bool]:
        """
        Extract text content from file
        Returns: (extracted_text, success)
        """
        try:
            extension = Path(filename).suffix.lower()
            
            if extension == '.txt':
                return cls._extract_from_txt(file_content)
            elif extension == '.pdf':
                return cls._extract_from_pdf(file_content)
            elif extension in ['.docx']:
                return cls._extract_from_docx(file_content)
            elif extension in ['.xlsx']:
                return cls._extract_from_xlsx(file_content)
            elif extension == '.csv':
                return cls._extract_from_csv(file_content)
            elif extension in ['.md', '.rtf']:
                return cls._extract_from_txt(file_content)  # Treat as plain text
            else:
                logging.warning(f"Unsupported file format: {extension}")
                return "", False
                
        except Exception as e:
            logging.error(f"Error extracting text from {filename}: {e}")
            return "", False
    
    @classmethod
    def _extract_from_txt(cls, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text, True
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='ignore')
            return text, True
            
        except Exception as e:
            logging.error(f"Error extracting text from TXT: {e}")
            return "", False
    
    @classmethod
    def _extract_from_pdf(cls, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from PDF files"""
        if not PDF_AVAILABLE:
            return "PDF processing not available. Please install PyPDF2.", False
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip(), True
            
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {e}")
            return "", False
    
    @classmethod
    def _extract_from_docx(cls, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            return "DOCX processing not available. Please install python-docx.", False
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip(), True
            
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {e}")
            return "", False
    
    @classmethod
    def _extract_from_xlsx(cls, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from XLSX files"""
        if not EXCEL_AVAILABLE:
            return "Excel processing not available. Please install openpyxl.", False
        
        try:
            xlsx_file = io.BytesIO(file_content)
            workbook = openpyxl.load_workbook(xlsx_file)
            
            text = ""
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            return text.strip(), True
            
        except Exception as e:
            logging.error(f"Error extracting text from XLSX: {e}")
            return "", False
    
    @classmethod
    def _extract_from_csv(cls, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from CSV files"""
        try:
            # Try different encodings for CSV
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    # Convert CSV to readable format
                    lines = text.split('\n')
                    formatted_text = ""
                    
                    for line in lines:
                        if line.strip():
                            # Replace commas with spaces for better readability
                            formatted_line = line.replace(',', ' | ')
                            formatted_text += formatted_line + "\n"
                    
                    return formatted_text.strip(), True
                except UnicodeDecodeError:
                    continue
            
            return "", False
            
        except Exception as e:
            logging.error(f"Error extracting text from CSV: {e}")
            return "", False